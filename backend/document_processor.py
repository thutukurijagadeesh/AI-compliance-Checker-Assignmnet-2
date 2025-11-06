import pdfplumber
from docx import Document
import spacy
from typing import Dict, List
from ml_model import ml_extractor, ml_risk_identifier
import os
import tempfile

nlp = spacy.load("en_core_web_sm")

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file."""
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() + "\n"
    return text

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file."""
    doc = Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def extract_key_clauses(text: str) -> List[str]:
    """Extract key clauses from contract text using NLP and ML."""
    doc = nlp(text)
    clauses = []

    # Look for sentences containing key terms
    key_terms = [
        "data processing", "personal data", "privacy", "consent", "breach",
        "security", "retention", "transfer", "rights", "liability",
        "termination", "confidentiality", "intellectual property",
        "agreement", "contract", "party", "parties", "hereby", "shall",
        "obligations", "terms", "conditions", "governing law"
    ]

    for sent in doc.sents:
        sent_text = sent.text.lower()
        if any(term in sent_text for term in key_terms):
            # Use ML to verify if it's a key clause
            ml_pred = ml_extractor.predict_clauses_ml(sent.text)
            if ml_pred['is_key_clause']:
                clauses.append(sent.text.strip())

    # If not enough clauses found, use GPT for additional extraction
    if len(clauses) < 5:
        gpt_clauses = extract_key_clauses_gpt(text)
        clauses.extend(gpt_clauses)
        clauses = list(set(clauses))  # Remove duplicates

    return clauses[:10]  # Limit to top 10

def extract_key_clauses_gpt(text: str) -> List[str]:
    """Extract key clauses using GPT."""
    # Fallback to rule-based extraction since GPT API is not available
    print("GPT not available, using rule-based clause extraction")
    return []

def identify_risks(text: str) -> List[Dict[str, str]]:
    """Identify compliance risks for GDPR and HIPAA using ML and GPT."""
    risks = []
    seen_risks = set()  # Track unique risk types

    # Use ML-based risk identification
    ml_risks = ml_risk_identifier.identify_risks_ml(text)
    for risk in ml_risks:
        risk_key = risk['risk'].lower()
        if risk_key not in seen_risks:
            risks.append(risk)
            seen_risks.add(risk_key)

    # Use GPT for additional risk identification
    gpt_risks = ml_risk_identifier.identify_risks_gpt(text)
    for gpt_risk in gpt_risks:
        risk_key = gpt_risk['risk'].lower()
        # Check if similar risk already exists
        if risk_key not in seen_risks and not any(r['risk'].lower() in risk_key or risk_key in r['risk'].lower() for r in risks):
            risks.append(gpt_risk)
            seen_risks.add(risk_key)

    # Fallback to rule-based if not enough risks found
    if len(risks) < 3:
        rule_based_risks = identify_risks_rule_based(text)
        for rb_risk in rule_based_risks:
            risk_key = rb_risk['risk'].lower()
            if risk_key not in seen_risks:
                risks.append(rb_risk)
                seen_risks.add(risk_key)

    # Remove duplicates by consolidating similar risks
    consolidated_risks = []
    risk_groups = {}

    for risk in risks:
        risk_key = risk['risk'].lower()
        # Group similar risks
        base_key = None
        for key in ['data processing', 'consent', 'breach', 'security', 'privacy', 'retention', 'transfer', 'rights']:
            if key in risk_key:
                base_key = key
                break

        if base_key:
            if base_key not in risk_groups:
                risk_groups[base_key] = risk
            else:
                # Keep the one with higher severity
                current_severity = risk_groups[base_key].get('severity', 'Low')
                new_severity = risk.get('severity', 'Low')
                severity_order = {'Low': 1, 'Medium': 2, 'High': 3}
                if severity_order.get(new_severity, 1) > severity_order.get(current_severity, 1):
                    risk_groups[base_key] = risk
        else:
            consolidated_risks.append(risk)

    # Add consolidated risks
    consolidated_risks.extend(risk_groups.values())

    return consolidated_risks[:10]  # Limit to top 10

def identify_risks_rule_based(text: str) -> List[Dict[str, str]]:
    """Rule-based risk identification as fallback."""
    risks = []

    # GDPR risks
    gdpr_risks = {
        "data processing without consent": "GDPR violation - processing personal data without lawful basis",
        "international data transfer": "GDPR risk - data transfers may require adequacy or safeguards",
        "data breach notification": "GDPR requirement - must notify breaches within 72 hours",
        "data subject rights": "GDPR requirement - must respond to data subject requests",
        "data retention": "GDPR risk - excessive data retention without justification"
    }

    # HIPAA risks
    hipaa_risks = {
        "phi disclosure": "HIPAA violation - unauthorized disclosure of protected health information",
        "security breach": "HIPAA requirement - must implement security safeguards",
        "patient rights": "HIPAA requirement - must provide access to medical records",
        "business associate agreement": "HIPAA requirement - contracts with service providers"
    }

    text_lower = text.lower()

    for risk_term, description in {**gdpr_risks, **hipaa_risks}.items():
        if risk_term.replace(" ", "") in text_lower.replace(" ", "") or any(word in text_lower for word in risk_term.split()):
            risks.append({
                "risk": risk_term,
                "description": description,
                "severity": "High" if "breach" in risk_term or "violation" in description else "Medium"
            })

    return risks

def identify_missing_clauses(text: str) -> List[Dict[str, str]]:
    """Identify missing clauses that are typically required in compliance documents."""
    text_lower = text.lower()
    missing_clauses = []

    # Required clauses for GDPR and HIPAA compliance
    required_clauses = {
        "data processing agreement": "GDPR requirement - outlines lawful processing of personal data",
        "consent clause": "GDPR/HIPAA requirement - explicit consent for data processing/use",
        "data security measures": "GDPR/HIPAA requirement - safeguards to protect data",
        "data retention policy": "GDPR requirement - specifies how long data is kept",
        "data subject rights": "GDPR requirement - rights of individuals over their data",
        "breach notification procedure": "GDPR/HIPAA requirement - process for notifying breaches",
        "international data transfers": "GDPR requirement - safeguards for cross-border transfers",
        "business associate agreement": "HIPAA requirement - contracts with service providers",
        "privacy notice": "GDPR/HIPAA requirement - information about data practices",
        "data minimization": "GDPR requirement - collect only necessary data"
    }

    for clause_term, description in required_clauses.items():
        # Check if the clause is mentioned in the text
        if not (clause_term.replace(" ", "") in text_lower.replace(" ", "") or any(word in text_lower for word in clause_term.split())):
            missing_clauses.append({
                "clause": clause_term,
                "description": description,
                "importance": "High" if "breach" in clause_term or "security" in clause_term else "Medium"
            })

    return missing_clauses[:10]  # Limit to top 10

def generate_recommended_clauses_gpt(document_type: str, risks: List[Dict], missing_clauses: List[Dict]) -> List[str]:
    """Generate recommended clauses using GPT based on document type, risks, and missing clauses."""
    from ml_model import client
    if client is None:
        return []

    # Prepare prompt with document type and risks
    prompt = f"""
    Based on the document type "{document_type}", analyze the following compliance risks and missing clauses.
    Generate 5-10 specific, legally sound recommended clauses that should be added to address these issues.

    High-Risk Issues:
    {chr(10).join([f"- {risk['risk']}: {risk['description']} (Severity: {risk.get('severity', 'Medium')})" for risk in risks if risk.get('severity') == 'High'])}

    Missing Clauses:
    {chr(10).join([f"- {clause['clause']}: {clause['description']}" for clause in missing_clauses])}

    Focus on GDPR, HIPAA, and general data protection requirements. Provide clauses that are:
    1. Specific to the document type
    2. Address the identified risks
    3. Fill the gaps from missing clauses
    4. Legally compliant and practical

    Format each clause as: "Clause Title: Full clause text"
    """

    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=1500,
            temperature=0.3
        )

        result = response.choices[0].message.content.strip()
        # Parse the response into a list of clauses
        clauses = [line.strip() for line in result.split('\n') if line.strip() and ':' in line]
        return clauses[:10]
    except Exception as e:
        print(f"GPT clause generation error: {e}")
        return []

def generate_recommended_clauses(document_type: str, risks: List[Dict], missing_clauses: List[Dict]) -> List[str]:
    """Generate recommended clauses based on document type, high-risk issues, and missing clauses using GPT or templates."""
    # Try GPT first
    gpt_clauses = generate_recommended_clauses_gpt(document_type, risks, missing_clauses)
    if gpt_clauses:
        return gpt_clauses

    # Fallback to template-based generation
    recommended_clauses = []

    # Templates for recommended clauses
    clause_templates = {
        "breach": [
            "Breach Notification Clause: In the event of a data breach, the data controller shall notify all affected data subjects within 72 hours of becoming aware of the breach, in accordance with GDPR Article 33.",
            "Breach Response Procedure: Upon discovery of a security breach, the parties agree to immediately implement their respective breach response procedures, including notification to supervisory authorities and affected individuals."
        ],
        "liability": [
            "Liability Limitation Clause: The total liability of either party under this agreement shall not exceed the total fees paid by the client in the twelve (12) months preceding the claim.",
            "Indemnification Clause: Each party agrees to indemnify and hold harmless the other party from any claims, damages, or liabilities arising from their negligence or willful misconduct."
        ],
        "data processing": [
            "Data Processing Agreement: The processor shall process personal data only on documented instructions from the controller, including with regard to transfers of personal data to a third country or an international organization.",
            "Lawful Basis Clause: Personal data shall be processed on a lawful basis, with explicit consent obtained where required by applicable data protection laws."
        ],
        "consent": [
            "Consent Clause: The controller shall obtain explicit, informed consent from data subjects before processing their personal data for any purpose beyond the original collection purpose.",
            "Withdrawal of Consent: Data subjects shall have the right to withdraw their consent at any time, and such withdrawal shall be as easy as giving consent."
        ],
        "security": [
            "Security Measures Clause: Both parties shall implement appropriate technical and organizational measures to ensure a level of security appropriate to the risk, including encryption, access controls, and regular security assessments.",
            "Data Protection Impact Assessment: Where processing is likely to result in high risk to the rights and freedoms of individuals, the controller shall carry out a data protection impact assessment."
        ],
        "retention": [
            "Data Retention Policy: Personal data shall be kept in a form which permits identification of data subjects for no longer than is necessary for the purposes for which the personal data are processed.",
            "Data Deletion Clause: Upon expiration or termination of this agreement, all personal data shall be securely deleted or returned to the controller within thirty (30) days."
        ],
        "rights": [
            "Data Subject Rights Clause: Data subjects shall have the right to access, rectify, erase, restrict processing, data portability, and object to processing of their personal data.",
            "Right to Information: The controller shall provide data subjects with information about the processing of their personal data in a concise, transparent, intelligible, and easily accessible form."
        ],
        "transfer": [
            "International Data Transfers Clause: Transfers of personal data to third countries shall only occur where the European Commission has decided that the third country ensures an adequate level of protection.",
            "Standard Contractual Clauses: Where adequate protection cannot be ensured, transfers shall be governed by standard contractual clauses approved by the European Commission."
        ],
        "privacy": [
            "Privacy Notice Clause: The controller shall provide data subjects with a privacy notice that includes information about the purposes of processing, legal basis, recipients of personal data, and data subject rights.",
            "Information Provision Clause: Data subjects shall be informed about the identity of the controller, purposes of processing, categories of personal data, and their rights under applicable data protection laws."
        ],
        "minimization": [
            "Data Minimization Clause: Personal data shall be adequate, relevant, and limited to what is necessary in relation to the purposes for which they are processed.",
            "Purpose Limitation Clause: Personal data shall be collected for specified, explicit, and legitimate purposes and not further processed in a manner that is incompatible with those purposes."
        ],
        "agreement": [
            "Business Associate Agreement: The business associate shall enter into a written agreement with the covered entity that establishes the permitted and required uses and disclosures of protected health information.",
            "Contractual Obligations Clause: The business associate shall comply with the requirements of the Privacy, Security, and Breach Notification Rules with respect to protected health information."
        ]
    }

    # Generate clauses for high-risk issues
    high_risks = [risk for risk in risks if risk.get('severity') == 'High']
    for risk in high_risks:
        risk_type = risk['risk'].lower()
        if 'breach' in risk_type:
            recommended_clauses.extend(clause_templates.get('breach', []))
        elif 'liability' in risk_type:
            recommended_clauses.extend(clause_templates.get('liability', []))
        elif 'processing' in risk_type or 'data processing' in risk_type:
            recommended_clauses.extend(clause_templates.get('data processing', []))
        elif 'consent' in risk_type:
            recommended_clauses.extend(clause_templates.get('consent', []))
        elif 'security' in risk_type:
            recommended_clauses.extend(clause_templates.get('security', []))
        elif 'retention' in risk_type:
            recommended_clauses.extend(clause_templates.get('retention', []))
        elif 'rights' in risk_type:
            recommended_clauses.extend(clause_templates.get('rights', []))
        elif 'transfer' in risk_type:
            recommended_clauses.extend(clause_templates.get('transfer', []))
        elif 'agreement' in risk_type:
            recommended_clauses.extend(clause_templates.get('agreement', []))

    # Generate clauses for missing clauses
    for missing in missing_clauses:
        clause_type = missing['clause'].lower()
        if 'breach' in clause_type:
            recommended_clauses.extend(clause_templates.get('breach', []))
        elif 'liability' in clause_type:
            recommended_clauses.extend(clause_templates.get('liability', []))
        elif 'processing' in clause_type or 'data processing' in clause_type:
            recommended_clauses.extend(clause_templates.get('data processing', []))
        elif 'consent' in clause_type:
            recommended_clauses.extend(clause_templates.get('consent', []))
        elif 'security' in clause_type:
            recommended_clauses.extend(clause_templates.get('security', []))
        elif 'retention' in clause_type:
            recommended_clauses.extend(clause_templates.get('retention', []))
        elif 'rights' in clause_type:
            recommended_clauses.extend(clause_templates.get('rights', []))
        elif 'transfer' in clause_type:
            recommended_clauses.extend(clause_templates.get('transfer', []))
        elif 'privacy' in clause_type or 'notice' in clause_type:
            recommended_clauses.extend(clause_templates.get('privacy', []))
        elif 'minimization' in clause_type:
            recommended_clauses.extend(clause_templates.get('minimization', []))
        elif 'agreement' in clause_type or 'business associate' in clause_type:
            recommended_clauses.extend(clause_templates.get('agreement', []))

    # Remove duplicates and limit to 10
    recommended_clauses = list(set(recommended_clauses))[:10]

    return recommended_clauses

def find_insertion_points(text: str, high_risks: List[Dict]) -> List[Dict]:
    """Find appropriate insertion points for high-risk issues based on content analysis."""
    paragraphs = text.split('\n')
    insertion_points = []

    for i, para in enumerate(paragraphs):
        para_lower = para.lower()
        for risk in high_risks:
            risk_type = risk['risk'].lower()
            # Find paragraphs that discuss related topics
            if any(keyword in para_lower for keyword in risk_type.split()):
                insertion_points.append({
                    'position': i + 1,  # Insert after this paragraph
                    'risk': risk
                })
                break  # Only one insertion per paragraph

    # If no specific insertion points found, add at the end
    if not insertion_points:
        for risk in high_risks:
            insertion_points.append({
                'position': len(paragraphs),
                'risk': risk
            })

    return insertion_points

def find_clause_insertion_points(text: str, recommended_clauses: List[str]) -> List[Dict]:
    """Find appropriate insertion points for recommended clauses based on content analysis."""
    paragraphs = text.split('\n')
    insertion_points = []

    for i, para in enumerate(paragraphs):
        para_lower = para.lower()
        for clause in recommended_clauses:
            clause_lower = clause.lower()
            # Look for sections that might be related (e.g., if clause mentions "privacy" and paragraph has privacy-related terms)
            if any(keyword in para_lower for keyword in ['privacy', 'data', 'security', 'consent', 'breach', 'liability']):
                if any(keyword in clause_lower for keyword in para_lower.split()):
                    insertion_points.append({
                        'position': i + 1,
                        'clause': clause
                    })
                    break

    # If no specific points found, distribute clauses throughout the document
    if not insertion_points:
        total_paragraphs = len(paragraphs)
        clauses_per_section = max(1, len(recommended_clauses) // 3)
        position = total_paragraphs // 4  # Start at 25% through document

        for clause in recommended_clauses:
            insertion_points.append({
                'position': min(position, total_paragraphs),
                'clause': clause
            })
            position += total_paragraphs // 4

    return insertion_points

def create_updated_document(original_file_path: str, file_type: str, text: str, high_risks: List[Dict], recommended_clauses: List[str]) -> str:
    """Create an updated document with high-risk items and recommended clauses inserted contextually."""
    # Create a temporary directory for updated files
    temp_dir = "temp"
    os.makedirs(temp_dir, exist_ok=True)

    # Generate filename for updated document
    base_name = os.path.splitext(os.path.basename(original_file_path))[0]
    updated_filename = f"{base_name}_updated.docx"
    updated_file_path = os.path.join(temp_dir, updated_filename)

    # Create new DOCX document
    doc = Document()

    # Add original content
    if file_type == "docx":
        original_doc = Document(original_file_path)
        for para in original_doc.paragraphs:
            doc.add_paragraph(para.text)
    else:
        # For PDF and TXT, add extracted text
        paragraphs = text.split('\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())

    # Insert high-risk items contextually - limit to unique risks
    if high_risks:
        # Remove duplicates from high_risks list
        unique_high_risks = []
        seen_risk_keys = set()
        for risk in high_risks:
            risk_key = risk['risk'].lower()
            if risk_key not in seen_risk_keys:
                unique_high_risks.append(risk)
                seen_risk_keys.add(risk_key)

        # Find appropriate insertion points based on content
        insertion_points = find_insertion_points(text, unique_high_risks)

        # Sort insertion points by position
        insertion_points.sort(key=lambda x: x['position'])

        # Insert high-risk sections - only once per unique risk
        inserted_risks = set()
        for point in insertion_points:
            risk_key = point["risk"]["risk"].lower()
            if risk_key not in inserted_risks:
                # Add a page break before high-risk section if not at the end
                if point['position'] < len(doc.paragraphs):
                    doc.paragraphs[point['position']].insert_paragraph_before('\n')

                doc.add_heading(f'High-Risk Compliance Issue: {point["risk"]["risk"]}', level=2)
                doc.add_paragraph(f'Description: {point["risk"]["description"]}')
                doc.add_paragraph(f'Severity: {point["risk"]["severity"]}')
                doc.add_paragraph(f'Recommended Action: Address this risk by adding appropriate clauses.')
                doc.add_paragraph()  # Empty paragraph for spacing
                inserted_risks.add(risk_key)

    # Insert recommended clauses contextually - limit to unique clauses
    if recommended_clauses:
        # Remove duplicates from recommended_clauses
        unique_clauses = list(set(recommended_clauses))

        # Find insertion points for recommended clauses
        clause_insertion_points = find_clause_insertion_points(text, unique_clauses)

        # Sort by position
        clause_insertion_points.sort(key=lambda x: x['position'])

        # Insert clauses - only once per unique clause
        inserted_clauses = set()
        for point in clause_insertion_points:
            clause_key = point["clause"].lower()
            if clause_key not in inserted_clauses:
                # Insert the clause with annotation
                clause_title = point["clause"].split(":")[0] if ":" in point["clause"] else "Compliance Clause"
                insertion_text = f'[ADDED CLAUSE - {clause_title}]\n{point["clause"]}'

                if point['position'] < len(doc.paragraphs):
                    doc.paragraphs[point['position']].insert_paragraph_before(insertion_text)
                else:
                    doc.add_paragraph(insertion_text)
                doc.add_paragraph()  # Empty paragraph for spacing
                inserted_clauses.add(clause_key)

    if not high_risks and not recommended_clauses:
        doc.add_page_break()
        doc.add_heading('Compliance Analysis Summary', level=1)
        doc.add_paragraph("No high-risk issues or specific recommendations generated based on the analysis.")

    # Save the updated document
    doc.save(updated_file_path)

    return updated_filename

def process_document(file_path: str, file_type: str) -> Dict:
    """Process uploaded document and return analysis."""
    try:
        if file_type == "pdf":
            text = extract_text_from_pdf(file_path)
        elif file_type == "docx":
            text = extract_text_from_docx(file_path)
        elif file_type == "txt":
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        else:
            raise ValueError("Unsupported file type")

        # First, classify if the document is a contract
        classification = ml_risk_identifier.classify_document_type(text)

        if not classification.get('is_contract', True):
            return {
                "is_contract": False,
                "message": "Please upload contract documents. This document does not appear to be a legal contract or agreement.",
                "key_clauses": [],
                "risks": [],
                "missing_clauses": [],
                "text_length": len(text),
                "analysis_method": "Document classification"
            }

        clauses = extract_key_clauses(text)

        # Additional check: if no key clauses found, likely not a contract
        if len(clauses) == 0:
            return {
                "is_contract": False,
                "message": "Please upload contract documents. This document does not appear to be a legal contract or agreement.",
                "key_clauses": [],
                "risks": [],
                "missing_clauses": [],
                "text_length": len(text),
                "analysis_method": "Document classification"
            }

        risks = identify_risks(text)
        missing_clauses = identify_missing_clauses(text)

        # Generate recommended clauses for high risks and missing clauses
        document_type = classification.get('document_type', 'General Contract')
        recommended_clauses = generate_recommended_clauses(document_type, risks, missing_clauses)

        # Get high-risk items
        high_risks = [risk for risk in risks if risk.get('severity') == 'High']

        # Create updated document if there are high risks or recommendations
        updated_filename = None
        if high_risks or recommended_clauses:
            updated_filename = create_updated_document(file_path, file_type, text, high_risks, recommended_clauses)

        return {
            "is_contract": True,
            "key_clauses": clauses,
            "risks": risks,
            "missing_clauses": missing_clauses,
            "recommended_clauses": recommended_clauses,
            "updated_filename": updated_filename,
            "text_length": len(text),
            "analysis_method": "ML-enhanced with GPT integration"
        }
    except Exception as e:
        return {
            "is_contract": False,
            "message": f"An error occurred during analysis: {str(e)}. Please try again or contact support.",
            "key_clauses": [],
            "risks": [],
            "missing_clauses": [],
            "text_length": 0,
            "analysis_method": "Error handling"
        }
